import os
import streamlit as st
import logging
from pathlib import Path
import sys
import time
import uuid
from docx import Document
from PIL import Image
import pytesseract
import io
import random

# Must be the first Streamlit command
st.set_page_config(page_title="Chatbot", layout="wide")

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import core modules
from core.utils import setup_api_keys, get_env_variable
from core.pdf_parser import PDFParser
from core.chunking import DocumentChunker
from embeddings.rag_engine import RAGEngine, RAGConfig
from core.ocr import OCRProcessor

# Initialize API keys
api_keys = setup_api_keys()

# Set default configuration
DEFAULT_MAX_TOKENS_CONTEXT = 8192
DEFAULT_MAX_TOKENS_RESPONSE = 2048
DEFAULT_MODEL = "meta-llama/Llama-4-Scout-17B-16E-Instruct"

# Load configuration from environment
max_tokens_context = int(get_env_variable("MAX_CONTEXT_TOKENS", DEFAULT_MAX_TOKENS_CONTEXT))
max_tokens_response = int(get_env_variable("MAX_RESPONSE_TOKENS", DEFAULT_MAX_TOKENS_RESPONSE))
model_name = get_env_variable("DEFAULT_MODEL", DEFAULT_MODEL)

# Create RAG configuration
rag_config = RAGConfig(
    llm_model=model_name,
    max_tokens_context=max_tokens_context,
    max_tokens_response=max_tokens_response,
    temperature=0.5,  # Reduced for more deterministic outputs
    retrieval_k=8,  # Increased for more comprehensive knowledge retrieval
    handle_low_quality=True  # Enable enhanced OCR and text correction by default
)

# Initialize session state for multiple chats
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = {}
    
if "current_chat_id" not in st.session_state:
    # Create initial chat
    initial_chat_id = str(uuid.uuid4())
    st.session_state.current_chat_id = initial_chat_id
    st.session_state.chat_sessions[initial_chat_id] = {
        "name": "New Chat",
        "messages": [],
        "rag_engine": None,
        "document_processed": False,
        "pdf_name": None,
        "direct_chat_engine": None
    }
    
# Add state for managing chat renaming
if "renaming_chat" not in st.session_state:
    st.session_state.renaming_chat = None
    
# Helper functions for chat management
def get_current_chat():
    """Get the current chat session data"""
    return st.session_state.chat_sessions[st.session_state.current_chat_id]

def set_current_chat(key, value):
    """Set a value in the current chat session"""
    st.session_state.chat_sessions[st.session_state.current_chat_id][key] = value
    
def create_new_chat():
    """Create a new chat session and switch to it"""
    new_chat_id = str(uuid.uuid4())
    st.session_state.chat_sessions[new_chat_id] = {
        "name": f"New Chat {len(st.session_state.chat_sessions)}",
        "messages": [],
        "rag_engine": None,
        "document_processed": False,
        "pdf_name": None,
        "direct_chat_engine": None
    }
    st.session_state.current_chat_id = new_chat_id
    
def switch_to_chat(chat_id):
    """Switch to the specified chat"""
    if chat_id in st.session_state.chat_sessions:
        st.session_state.current_chat_id = chat_id
        
def clear_conversation():
    """Clear the conversation history in the current chat"""
    set_current_chat("messages", [])
    
def clear_document():
    """Clear the document in the current chat"""
    set_current_chat("document_processed", False)
    set_current_chat("rag_engine", None)
    set_current_chat("pdf_name", None)
    
def regenerate_response():
    """Regenerate the last response"""
    current_chat = get_current_chat()
    messages = current_chat["messages"]
    
    # Need at least a user message to regenerate
    if len(messages) < 1:
        return
    
    # Find the last user message
    last_user_msg_idx = None
    for i in range(len(messages)-1, -1, -1):
        if messages[i]["role"] == "user":
            last_user_msg_idx = i
            break
            
    if last_user_msg_idx is not None:
        # Remove all messages after the last user message
        last_user_msg = messages[last_user_msg_idx]
        set_current_chat("messages", messages[:last_user_msg_idx+1])
        # Reprocess the query
        process_query(last_user_msg["content"])

# Sidebar for file upload, settings, and chat selection
with st.sidebar:
    st.title("You can create new chats from below ⬇️")
    
    # Chat selection section
    st.markdown("### Chats")
    
    # Display the list of chats
    for chat_id, chat_data in st.session_state.chat_sessions.items():
        chat_name = chat_data["name"]
        doc_name = chat_data["pdf_name"] if chat_data["pdf_name"] else "No document"
        
        # Create a row for each chat with three columns
        col1, col2, col3 = st.columns([3, 1, 1])
        
        with col1:
            # If this chat is being renamed, show text input
            if st.session_state.renaming_chat == chat_id:
                new_name = st.text_input("New name", value=chat_name, key=f"rename_{chat_id}")
                if new_name and new_name != chat_name:
                    st.session_state.chat_sessions[chat_id]["name"] = new_name
                    st.session_state.renaming_chat = None
                    st.rerun()
            else:
                # Highlight the active chat
                if chat_id == st.session_state.current_chat_id:
                    st.markdown(f"**→ {chat_name}** ({doc_name})")
                else:
                    if st.button(f"{chat_name} ({doc_name})", key=f"chat_{chat_id}"):
                        switch_to_chat(chat_id)
                        st.rerun()
        
        with col2:
            # Rename button for non-active chats
            if chat_id != st.session_state.current_chat_id:
                if st.button("✏️", key=f"rename_btn_{chat_id}"):
                    st.session_state.renaming_chat = chat_id
                    st.rerun()
        
        with col3:
            # Delete button for non-active chats when there are multiple chats
            if chat_id != st.session_state.current_chat_id and len(st.session_state.chat_sessions) > 1:
                if st.button("×", key=f"del_{chat_id}"):
                    del st.session_state.chat_sessions[chat_id]
                    st.rerun()
    
    # New chat button
    if st.button("+ New Chat"):
        create_new_chat()
        st.rerun()
    
    st.markdown("---")
    
    # Document upload section
    st.markdown("### Upload Document")
    current_chat = get_current_chat()
    uploaded_file = st.file_uploader(
        "Upload a document", 
        type=['pdf', 'jpg', 'jpeg', 'png', 'svg', 'docx'],
        key=f"uploader_{st.session_state.current_chat_id}"
    )
    
    # Advanced settings
    with st.expander("Advanced Settings"):
        context_tokens = st.slider(
            "Max Context Tokens", 
            min_value=1000, 
            max_value=8000, 
            value=rag_config.max_tokens_context,
            step=500,
            help="Maximum tokens for context (higher values may cause token limit errors)"
        )
        rag_config.max_tokens_context = context_tokens
        
        response_tokens = st.slider(
            "Max Response Tokens", 
            min_value=500, 
            max_value=2000, 
            value=rag_config.max_tokens_response,
            step=100,
            help="Maximum tokens for AI response"
        )
        rag_config.max_tokens_response = response_tokens
        
        temperature = st.slider(
            "Temperature", 
            min_value=0.0, 
            max_value=1.0, 
            value=rag_config.temperature,
            step=0.1,
            help="Controls randomness (0=deterministic, 1=creative)"
        )
        rag_config.temperature = temperature
        
        # Document processing settings
        st.markdown("### Document Processing")
        
        enhanced_ocr = st.toggle(
            "Enhanced OCR & Text Correction", 
            value=rag_config.handle_low_quality,
            help="Enable advanced OCR and statistical text correction for problematic PDFs"
        )
        rag_config.handle_low_quality = enhanced_ocr
        
        retrieval_k = st.slider(
            "Retrieved Chunks", 
            min_value=2, 
            max_value=8, 
            value=rag_config.retrieval_k,
            step=1,
            help="Number of document chunks to retrieve per query"
        )
        rag_config.retrieval_k = retrieval_k

    # API key status
    # st.markdown("### API Status")
    # if api_keys["huggingface_key"]:
    #     st.success("HuggingFace API key loaded")
    # else:
    #     st.warning("HuggingFace API key not found")
    #         
    # if api_keys["groq_key"]:
    #     st.success("Groq API key loaded")
    # else:
    #     st.error("Groq API key missing")

# Add file type validation function
def is_valid_file_type(file):
    """Check if the file is of a supported type."""
    valid_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.svg', '.docx']
    file_extension = os.path.splitext(file.name)[1].lower()
    return file_extension in valid_extensions

# Identity management function
def is_identity_question(query: str) -> bool:
    """
    Check if the query is asking about the bot's identity/creator.
    Returns True if it's an identity question, False otherwise.
    """
    # Convert to lowercase for case-insensitive matching
    query_lower = query.lower()
    
    # Keywords related to identity questions
    identity_keywords = [
        "who built you", "who made you", "who created you", 
        "who developed you", "who programmed you", "who designed you",
        "who owns you", "who are you made by", "who's your creator",
        "who is your creator", "who is your developer", "who is your maker",
        "are you made by", "are you created by", "are you developed by",
        "are you from", "where are you from", "what company made you",
        "which company", "what organization", "which organization",
        "who created this bot", "who built this bot", "who made this chatbot",
        "openai", "meta", "google", "microsoft", "anthropic", "claude",
        "moksha", "moksha solutions", "what is moksha", "what's moksha",
        "are you a real bot", "are you an ai"
    ]
    
    # Check if any of the keywords are in the query
    for keyword in identity_keywords:
        if keyword in query_lower:
            return True
    
    return False

# Friendly identity responses
def get_identity_response() -> str:
    """
    Returns a randomly selected friendly response about the bot's identity.
    """
    friendly_responses = [
        "Oh, I was built by the AI development team at Moksha Solutions! They're working on making document understanding smarter and more helpful for students with their studies and school work.",
        
        "The AI development team at Moksha Solutions is behind me. They've been experimenting with AI to build useful bots to help students learn more effectively. Pretty cool, right?",
        
        "That would be the AI development team at Moksha Solutions — the brains behind this chatbot! They're on a mission to make studying and research easier for students.",
        
        "I'm a project from Moksha Solutions. Their AI team designed me to help students chat with documents and get quick answers to support their studies without all the reading. How can I help you today?",
        
        "Moksha Solutions' AI team created me for educational purposes. I'm here to make studying and learning more intuitive and helpful for students!",
        
        "I'm one of Moksha Solutions' projects! Their AI development team built me to help students access and understand information more easily for their schoolwork.",
        
        "Yep, I'm a chatbot developed by the AI development team at Moksha Solutions. I specialize in helping students understand documents better and faster for their studies. What can I help you with?"
    ]
    
    return random.choice(friendly_responses)

# Add file processing function
def process_file(file):
    """Process different file types and extract text."""
    file_extension = os.path.splitext(file.name)[1].lower()
    
    try:
        if file_extension == '.pdf':
            # Use existing PDF processing
            text, metadata, doc_info = PDFParser().process_pdf(file.getvalue())
            return text, metadata, doc_info
            
        elif file_extension in ['.jpg', '.jpeg', '.png']:
            # Process image files
            image = Image.open(io.BytesIO(file.getvalue()))
            text = pytesseract.image_to_string(image)
            metadata = [{"text": text, "page": 1, "extraction_method": "ocr"}]
            doc_info = {"page_count": 1, "file_type": "image"}
            return text, metadata, doc_info
            
        elif file_extension == '.docx':
            # Process Word documents
            doc = Document(io.BytesIO(file.getvalue()))
            full_text = []
            metadata = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
                    metadata.append({
                        "text": paragraph.text,
                        "page": i + 1,
                        "extraction_method": "docx"
                    })
            
            text = "\n\n".join(full_text)
            doc_info = {"page_count": len(doc.paragraphs), "file_type": "docx"}
            return text, metadata, doc_info
            
        elif file_extension == '.svg':
            # Process SVG files (convert to PNG first)
            image = Image.open(io.BytesIO(file.getvalue()))
            # Convert to PNG for OCR
            png_buffer = io.BytesIO()
            image.save(png_buffer, format='PNG')
            png_buffer.seek(0)
            text = pytesseract.image_to_string(Image.open(png_buffer))
            metadata = [{"text": text, "page": 1, "extraction_method": "ocr"}]
            doc_info = {"page_count": 1, "file_type": "svg"}
            return text, metadata, doc_info
            
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        raise ValueError(f"Error processing file: {str(e)}")

# Update the process_pdf function to handle all file types
def process_document(file):
    """Process document and create knowledge base."""
    try:
        # Validate file type
        if not is_valid_file_type(file):
            raise ValueError("Unsupported file type. Please upload a PDF, JPG, PNG, SVG, or DOCX file.")
        
        # Get file bytes
        file_bytes = file.getvalue()
        
        # Add file size info
        file_size_mb = len(file_bytes) / (1024 * 1024)
        if file_size_mb > 10:
            st.warning(f"Large document detected ({file_size_mb:.1f} MB). Processing may take longer. Consider using a smaller document for faster results.")
        
        with st.status("Processing document...", expanded=True) as status:
            # Initialize progress bar
            progress_bar = st.progress(0)
            status.update(label="Analyzing document...", state="running")
            progress_bar.progress(10)
            
            # Process the file based on its type
            text, metadata, doc_info = process_file(file)
            progress_bar.progress(40)
            
            if not text:
                progress_bar.empty()
                st.error("No text could be extracted from the document. Please try a different file or enable Enhanced OCR in settings.")
                return False
                
            # Display document info
            status.update(label="Processing extracted text...", state="running")
            progress_bar.progress(50)
            
            st.write(f"Document: {file.name}")
            st.write(f"Type: {doc_info.get('file_type', 'Unknown')}")
            if 'page_count' in doc_info:
                st.write(f"Pages: {doc_info.get('page_count', 'Unknown')}")
            
            # Check for potential encoding issues in the document
            problematic_blocks = [block for block in metadata if block.get("corrected", False)]
            problematic_ratio = len(problematic_blocks) / max(len(metadata), 1)
            
            if problematic_ratio > 0.3:
                st.warning(f"Document appears to have encoding issues or custom fonts. Statistical text correction has been applied to improve readability.")
            elif problematic_blocks:
                st.info(f"Some parts of the document have been processed with text correction to improve quality.")
            
            # Check for large documents and warn user
            page_count = doc_info.get('page_count', 0)
            if page_count > 100:
                st.warning(f"Large document detected ({page_count} pages). Only processing a subset of pages for performance reasons.")
            
            # Create document chunks
            status.update(label="Creating document chunks...", state="running")
            progress_bar.progress(60)
            
            chunker = DocumentChunker(
                chunk_size=400,
                chunk_overlap=50,
                respect_sections=True
            )
            
            try:
                chunks = chunker.chunk_document(text, metadata)
                progress_bar.progress(70)
                
                # Merge small chunks if needed
                if chunks:
                    chunks = chunker.merge_small_chunks(chunks)
                    st.write(f"Created {len(chunks)} document chunks")
                else:
                    st.warning("Created empty chunks. The document might have formatting issues.")
                    # Create simpler chunks in case of failure
                    simple_chunks = []
                    lines = text.split('\n')
                    current_chunk = ""
                    
                    for line in lines:
                        if len(current_chunk) > 400:
                            simple_chunks.append({
                                "text": current_chunk,
                                "metadata": [],
                                "tokens": len(current_chunk) // 4
                            })
                            current_chunk = line
                        else:
                            current_chunk += line + "\n"
                    
                    if current_chunk:
                        simple_chunks.append({
                            "text": current_chunk,
                            "metadata": [],
                            "tokens": len(current_chunk) // 4
                        })
                    
                    chunks = simple_chunks
                    st.write(f"Created {len(chunks)} simple document chunks")
            except Exception as e:
                progress_bar.empty()
                st.error(f"Error creating chunks: {e}")
                return False
            
            # Initialize RAG engine
            status.update(label="Initializing RAG engine...", state="running")
            progress_bar.progress(80)
            
            try:
                rag_engine = RAGEngine(config=rag_config).initialize()
            except Exception as e:
                progress_bar.empty()
                st.error(f"Error initializing RAG engine: {e}")
                return False
            
            # Create knowledge base
            status.update(label="Building knowledge base...", state="running")
            progress_bar.progress(90)
            
            try:
                rag_engine.create_knowledge_base(chunks)
            except Exception as e:
                progress_bar.empty()
                st.error(f"Error creating knowledge base: {e}")
                return False
            
            # Setup retrieval chain
            status.update(label="Setting up AI model...", state="running")
            progress_bar.progress(95)
            
            try:
                groq_api_key = get_env_variable("GROQ_API_KEY")
                rag_engine.setup_retrieval_chain(groq_api_key)
            except Exception as e:
                progress_bar.empty()
                st.error(f"Error setting up retrieval chain: {e}")
                return False
            
            # Store in current chat session
            set_current_chat("rag_engine", rag_engine)
            set_current_chat("document_processed", True)
            set_current_chat("pdf_name", file.name)
            
            # Complete progress
            progress_bar.progress(100)
            status.update(label="Document processed successfully!", state="complete", expanded=False)
            
            # Clean up progress bar after a delay
            time.sleep(1)
            progress_bar.empty()
            
            return True
            
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        st.error(f"Error processing document: {str(e)}")
        return False

# Process PDF when uploaded
current_chat = get_current_chat()
if uploaded_file and not current_chat["document_processed"]:
    success = process_document(uploaded_file)
    
    if success:
        # Add a system message to the chat
        set_current_chat("messages", current_chat["messages"] + [{
            "role": "assistant", 
            "content": f"Document '{uploaded_file.name}' processed successfully! Ask me anything about it."
        }])
        st.rerun()

# Main chat interface
st.markdown("<h1 style='text-align: center;'>PaperTrail AI</h1>", unsafe_allow_html=True)
st.markdown("<p style='color: #808080; text-align: center; margin-bottom: 15px; font-size: 18px;'>What can I help with? Upload a document and chat with it</p>", unsafe_allow_html=True)
# update the title to be more descriptive

# Display document info if loaded
if current_chat["document_processed"]:
    st.write(f"Current document: {current_chat['pdf_name']}")
    
# Display chat messages
for message in current_chat["messages"]:
    with st.chat_message(message["role"]):
        st.write(message["content"])

# User input

query = st.chat_input("Ask anything about your document...")

# Function to process query
def process_query(query):
    # Get current chat data
    current_chat = get_current_chat()
    
    # Add user message to chat history
    new_messages = current_chat["messages"] + [{"role": "user", "content": query}]
    set_current_chat("messages", new_messages)
    
    # Check if this is an identity question
    if is_identity_question(query):
        with st.chat_message("assistant"):
            # Get a friendly identity response
            identity_response = get_identity_response()
            
            # Use the response
            st.write(identity_response)
            
            # Add the same response to chat history
            new_messages = current_chat["messages"] + [{"role": "assistant", "content": identity_response}]
            set_current_chat("messages", new_messages)
        return
    
    # Process with assistant
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            try:
                # Check if we have a document processed
                if current_chat["document_processed"] and current_chat["rag_engine"]:
                    # Use RAG-based response
                    answer, sources = current_chat["rag_engine"].query(query)
                    
                    # Display answer
                    st.write(answer)
                    
                    # Show sources if available
                    if sources:
                        with st.expander("Sources"):
                            for i, source in enumerate(sources):
                                st.markdown(f"**Source {i+1}:**")
                                
                                # Display page number if available
                                if "metadata" in source and "page" in source["metadata"]:
                                    st.markdown(f"Page: {source['metadata']['page']}")
                                
                                # Show if this source was corrected
                                if source.get("metadata", {}).get("corrected", False):
                                    st.markdown("*Text has been enhanced with statistical correction*")
                                
                                # Show content preview
                                content = source.get("content", "")
                                preview = content[:200] + "..." if len(content) > 200 else content
                                st.markdown(f"```\n{preview}\n```")
                                
                                # Show confidence if available
                                confidence = source.get("metadata", {}).get("confidence", None)
                                if confidence is not None:
                                    confidence_color = "green" if confidence > 0.8 else "orange" if confidence > 0.5 else "red"
                                    st.markdown(f"<span style='color:{confidence_color}'>Source confidence: {confidence:.2f}</span>", unsafe_allow_html=True)
                    
                    # Add assistant response to chat history
                    new_messages = current_chat["messages"] + [{"role": "assistant", "content": answer}]
                    set_current_chat("messages", new_messages)
                else:
                    # No document uploaded, use direct LLM chat instead
                    groq_api_key = get_env_variable("GROQ_API_KEY")
                    if not groq_api_key:
                        st.error("Groq API key is missing. Please add it to your environment variables.")
                        answer = "I couldn't process your request because the Groq API key is missing. Please add it to your environment variables."
                    else:
                        # Check if we have a direct chat engine already
                        if "direct_chat_engine" not in current_chat or current_chat["direct_chat_engine"] is None:
                            # Create and store a chat engine for this session
                            temp_engine = RAGEngine(config=rag_config).initialize()
                            temp_engine.setup_direct_chain(groq_api_key)
                            set_current_chat("direct_chat_engine", temp_engine)
                            
                        # Get the chat engine and query
                        chat_engine = current_chat["direct_chat_engine"]
                        answer = chat_engine.query_direct(query)
                    
                    # Display the answer
                    st.write(answer)
                    
                    # Add assistant response to chat history for direct chat
                    new_messages = current_chat["messages"] + [{"role": "assistant", "content": answer}]
                    set_current_chat("messages", new_messages)
                
            except Exception as e:
                error_msg = str(e)
                logger.error(f"Error generating response: {error_msg}")
                
                # Check if it's a token limit error
                if "token" in error_msg.lower() and "limit" in error_msg.lower():
                    recovery_msg = (
                        "I encountered a token limit error. This happens when the conversation "
                        "gets too long or the document sections needed are too large. "
                        "Try asking a more specific question or reducing the context tokens in settings."
                    )
                    st.error(error_msg)
                    st.write(recovery_msg)
                    new_messages = current_chat["messages"] + [{"role": "assistant", "content": recovery_msg}]
                    set_current_chat("messages", new_messages)
                else:
                    # General error
                    st.error(f"Error: {e}")
                    error_response = f"I encountered an error while processing your question. Please try again or ask a different question."
                    st.write(error_response)
                    new_messages = current_chat["messages"] + [{"role": "assistant", "content": error_response}]
                    set_current_chat("messages", new_messages)

# Process user query
if query:
    # Add user message to chat history
    with st.chat_message("user"):
        st.write(query)
    
    # Process the query
    process_query(query)
    
# UI controls section - at the bottom of the conversation
if current_chat["document_processed"]:
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("Clear Conversation"):
            clear_conversation()
            st.rerun()
            
    with col2:
        if st.button("Clear Document"):
            clear_document()
            clear_conversation()
            st.rerun()
            
    with col3:
        if st.button("Regenerate Response"):
            regenerate_response()
            st.rerun()
   